
import time
from io import BytesIO

from django.conf import settings
from django.http import HttpResponse
import os
from random import randint
from docx import Document
from docx.shared import Inches
import json


class Render:

    @staticmethod
    def render(nickname=str, params=dict):

        document = Document()
        response = json.loads(params.content)
        person = response['person']
        jobs = response['jobs']

        document.add_heading(person['name'], 0)

        document.add_paragraph(person['summaryOfBio'])

        document.add_heading('Experiencia', level=1)
        for job in jobs:
            document.add_paragraph(job['name'], style='Intense Quote')

            document.add_paragraph(job['additionalInfo'])
            for organization in job['organizations']:
                document.add_paragraph(organization['name'], style='List Bullet')

        # document.add_picture(person['pictureThumbnail'], width=Inches(1.25))

        document.add_heading('Contacto', level=1)
        for link in person['links']:
            document.add_paragraph('{}: {}'.format(link['name'], link['address']), style='List Bullet')

        file_name = "cv-{0}-{1}.docx".format(nickname, time.strftime("%Y%m%d-%H%M%S"))
        if not os.path.exists(os.path.join(settings.MEDIA_ROOT, "cvs")):
            os.mkdir(os.path.join(settings.MEDIA_ROOT, "cvs"))
        file_path = os.path.join(settings.MEDIA_ROOT, "cvs", file_name)

        document.save(file_path)

        return [file_name, file_path]
